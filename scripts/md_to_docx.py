#!/usr/bin/env python3
"""Convert Markdown to Word (.docx) using python-docx."""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def strip_frontmatter(text: str) -> tuple[dict, str]:
    """Strip YAML frontmatter and return (meta, body)."""
    meta = {}
    if not text.startswith('---'):
        return meta, text
    end = text.find('\n---', 3)
    if end == -1:
        return meta, text
    fm = text[4:end]
    body = text[end + 4:].lstrip('\n')
    for line in fm.splitlines():
        if ':' in line:
            k, _, v = line.partition(':')
            meta[k.strip()] = v.strip().strip('"')
    return meta, body


def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_hyperlink_run(para, text):
    """Add plain text run (no actual hyperlink needed for Word export)."""
    run = para.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x78, 0xD1)
    run.font.underline = True
    return run


def apply_inline(para, text: str):
    """Parse inline markdown (bold, italic, code, plain) into runs."""
    # Patterns: **bold**, *italic*, `code`, plain
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`)')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        if m.group(2):  # bold
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):  # italic
            run = para.add_run(m.group(3))
            run.italic = True
        elif m.group(4):  # code
            run = para.add_run(m.group(4))
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


def parse_table(lines: list[str], doc: Document):
    """Parse a markdown table and add to document."""
    rows = []
    for line in lines:
        if re.match(r'^\|[-:\s|]+\|?$', line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            cell = table.cell(ri, ci)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ri == 0 else WD_ALIGN_PARAGRAPH.LEFT
            apply_inline(para, cell_text)
            if ri == 0:
                for run in para.runs:
                    run.bold = True
                set_cell_bg(cell, 'D6E4F7')
    doc.add_paragraph()


def convert(md_path: str, out_path: str):
    text = Path(md_path).read_text(encoding='utf-8')
    meta, body = strip_frontmatter(text)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)

    # Default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    lines = body.splitlines()
    i = 0
    in_code_block = False
    code_lines = []
    table_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                in_code_block = False
                p = doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                doc.add_paragraph()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table accumulation
        if line.startswith('|'):
            table_lines.append(line)
            i += 1
            continue
        else:
            if table_lines:
                parse_table(table_lines, doc)
                table_lines = []

        # Headings
        heading_m = re.match(r'^(#{1,6})\s+(.*)', line)
        if heading_m:
            level = len(heading_m.group(1))
            text_h = heading_m.group(2)
            p = doc.add_heading(level=min(level, 4))
            p.clear()
            run = p.add_run(text_h)
            sizes = {1: 18, 2: 15, 3: 13, 4: 11}
            run.font.size = Pt(sizes.get(level, 11))
            run.font.bold = True
            run.font.name = '黑体'
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x6B)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Bullet list
        if re.match(r'^(\s*)([-*+])\s+(.*)', line):
            m = re.match(r'^(\s*)([-*+])\s+(.*)', line)
            indent = len(m.group(1)) // 2
            p = doc.add_paragraph(style='List Bullet')
            apply_inline(p, m.group(3))
            i += 1
            continue

        # Numbered list
        if re.match(r'^\s*\d+\.\s+(.*)', line):
            m = re.match(r'^\s*\d+\.\s+(.*)', line)
            p = doc.add_paragraph(style='List Number')
            apply_inline(p, m.group(1))
            i += 1
            continue

        # Blank line
        if line.strip() == '':
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        apply_inline(p, line.strip())
        i += 1

    # flush remaining table
    if table_lines:
        parse_table(table_lines, doc)

    doc.save(out_path)
    print(f'Saved: {out_path}')


if __name__ == '__main__':
    src = sys.argv[1] if len(sys.argv) > 1 else 'docs/plans/智慧水务平台需求说明书.md'
    dst = sys.argv[2] if len(sys.argv) > 2 else src.replace('.md', '.docx')
    convert(src, dst)
